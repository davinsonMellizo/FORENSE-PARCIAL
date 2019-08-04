#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import sys 
import subprocess 
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import Pt
from os import listdir
from os.path import isfile, join

def ls(ruta):
    return [arch for arch in listdir(ruta) if isfile(join(ruta, arch))]


def generarWord(info):

    insInicial=time.time()
    archivos = info.get('archivos')

    document = Document()
    l0 = document.add_heading('FORMATO DE CADENA DE CUSTODIA', 0)
    l0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    l1 = document.add_paragraph()
    l1.add_run('Numero de caso: ').bold = True
    l1.add_run(info.get('numCaso')+"\n")
    l1.add_run('Delito: ').bold = True
    l1.add_run(info.get('delito')+"\n")
    l1.add_run('Oficial encargado (Nombre/id): ').bold = True
    l1.add_run(info.get('nomOfEncargado')+"\t"+info.get('idOfEncargado')+"\n")
    l1.add_run('Nombre de la victima: ').bold = True
    l1.add_run(info.get('nomVictima')+"\n")
    l1.add_run('Nombre del sospechoso: ').bold = True
    l1.add_run(info.get('nomSospechoso')+"\n")
    l1.add_run('Nombre de la victima: ').bold = True
    l1.add_run(info.get('nomVictima')+"\n")
    l1.add_run('Fecha de la incuatacion: ').bold = True
    l1.add_run(info.get('fecha')+"\n")
    l1.add_run('Lugar de la incuatacion: ').bold = True
    l1.add_run(str(info.get('lugar'))+"\n")

    t1 = document.add_paragraph()
    t1.add_run("DESCRIPCION DE LA EVIDENCIA").bold = True
    t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    table = document.add_table(rows=1, cols=4)
    table.style = 'Light List Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#Item'
    hdr_cells[0].width
    hdr_cells[1].text = 'Nombre del item'
    hdr_cells[2].text = 'Descripcion'
    hdr_cells[3].text = 'Hash'

    for item in archivos:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('numItem')
        row_cells[1].text = item.get('nomArchivo')
        row_cells[2].text = item.get('desArchivo')
        row_cells[3].text = item.get('hash')


    t2 = document.add_paragraph()
    t2.add_run("\nCADENA DE CUSTODIA").bold = True
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = document.add_table(rows=2, cols=5)
    table1.style = 'Light List Accent 1'

    hdr1_cells = table1.rows[0].cells
    hdr1_cells[0].text = '#Item'
    hdr1_cells[1].text = 'Fecha/hora'
    hdr1_cells[2].text = 'Publicado por (Firma y #ID)'
    hdr1_cells[3].text = 'Resivido por (Firma y #ID)'
    hdr1_cells[4].text = 'Comentarios / Ubicacion'

    t3 = document.add_paragraph("\n")  

    table2 = document.add_table(rows=5, cols=1)
    table2.style = 'Light List Accent 1'

    hdr2_cells = table2.rows[0].cells
    tx = hdr2_cells[0].add_paragraph()
    tx.add_run('Autoridad de Disposicion Final\n').bold = True
    tx.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr2_cells = table2.rows[1].cells
    tx0 = hdr2_cells[0].add_paragraph()
    tx0.add_run('Autorizacion para la eliminacion\n').bold = True
    tx0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tx1 = hdr2_cells[0].add_paragraph()
    tx1.add_run('Articulo (s) #: __________ en este documento que pertenece a (sospechoso): ____________________________________________ ya no se necesita mas como evidencia y esta autorizado para su eliminacion por (verifique el metodo de eliminacion apropiado).')
    tx1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tx2 = hdr2_cells[0].add_paragraph()
    tx2.add_run('__ Volver al propietario \t\t___ Subasta / Destruir / Desviar')
    tx2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tx3 = hdr2_cells[0].add_paragraph()
    tx3.add_run('Nombre y numero de identificacion del oficial de autorizacion: ____________________________ Firma: ______________________ Fecha: _______________')
    tx3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    hdr2_cells = table2.rows[2].cells 
    tx0 = hdr2_cells[0].add_paragraph()
    tx0.add_run('Testigo de la destruccion de la evidencia\n').bold = True
    tx0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tx1 = hdr2_cells[0].add_paragraph()
    tx1.add_run('Articulo (s) #: __________ en este documento fueron destruidos por el Custodio de Evidencia ___________________________ Identificacion #: ______ en mi presencia en la (fecha) __________________________.')
    tx1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tx2 = hdr2_cells[0].add_paragraph()
    tx2.add_run('Nombre y numero de identificacion del testigo de la destruccion: ________________________ \nFirma: ______________________ Fecha: _______________')
    tx2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    hdr2_cells = table2.rows[3].cells 
    tx0 = hdr2_cells[0].add_paragraph()
    tx0.add_run('Liberacion al propietario legal\n').bold = True
    tx0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tx1 = hdr2_cells[0].add_paragraph()
    tx1.add_run('El articulo (s) #: __________ en este documento fue / fue publicado por el Custodio de Evidencia ________________________ # de ID: _________ para \nNombre _____________________________________________________________________________\nDireccion: ________________________________________________ Ciudad: ____________________ Estado: _______ Codigo postal: __________\nNumero de telefono: (_____) ___________________________________ Bajo la pena de ley, certifico que soy el propietario legitimo del (los) articulo (s) anterior (es).\nFirma:_______________________________________________________\nFecha: __________________________')
    tx1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    tx2 = hdr2_cells[0].add_paragraph()
    tx2.add_run('Se adjunta una copia de la identificacion con foto emitida por el gobierno. __ Si __No')
    tx2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    hdr2_cells = table2.rows[4].cells 
    tx0 = hdr2_cells[0].add_paragraph()
    tx0.add_run('Este formulario de Cadena de Custodia de Evidencia debe conservarse como un registro permanente por el Departamento de Policia de cualquier lugar.').bold = True
    tx0.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.save('Informe_Forense_casos_num_'+info.get('numCaso')+'.docx')
    print('Archivo generado: '+'Informe_Forense_Caso_Num_'+info.get('numCaso')+'.docx')
    insFinal=time.time()
    tiempo=insFinal- insInicial
    print("Tiempo de Ejecucion", tiempo)


def cadenaCustodia(ruta):

    
    ult = len(ruta)-1

    if ruta[ult] != "/":
        ruta = ruta+"/"

    listaArchivos = ls(ruta)
    numEvidencias = len(listaArchivos)
        
    if numEvidencias != 0:
        
        info = {}
        
        print("------FORULARIO DE CADENA DE CUSTODIA------")
        info['numCaso'] = raw_input("Digite el número del caso: ")
        info['delito'] = raw_input("Digite el tipo de delito: ")
        print("---------DATOS DEL OFICIAL ENCARGADO-------")
        info['nomOfEncargado'] = raw_input("Nombre: ")
        info['idOfEncargado'] = raw_input("Identificación: ")
        print("-----------DATOS DE LA VICTIMA-------------")
        info['nomVictima'] = raw_input("Nombre: ")
        print("----------DATOS DEL SOSPECHOSO-------------")
        info['nomSospechoso'] = raw_input("Nombre: ")
        print("---------DATOS DE LA INCUATACION-----------")
        info['fecha'] = time.strftime("%c")
        info['lugar'] = raw_input("Ingrese lugar de la incuatación: ")
        print("Generando archivo......")
        
        i = 0
        listaItems = []
        while (i<numEvidencias):
            item = {}
            item['numItem'] = str(i+1)
            item['nomArchivo'] = listaArchivos[i]
            item['desArchivo'] = subprocess.check_output('file '+ruta+listaArchivos[i],shell=True)
            item['hash'] = subprocess.check_output('md5sum '+ruta+listaArchivos[i],shell=True)[0:31]     
            listaItems.append(item)
            i = i+1
        info['archivos'] = listaItems
        generarWord(info)    

    else:
        print("Directori vacio")


def cadenaCustodiaVacia(ruta):

    ult = len(ruta)-1

    if ruta[ult] != "/":
        ruta = ruta+"/"
    
    listaArchivos = ls(ruta)
    numEvidencias = len(listaArchivos)
        
    if numEvidencias != 0:
        
        info = {}
        
        info['numCaso'] = ""
        info['delito'] = ""
        info['nomOfEncargado'] = ""
        info['idOfEncargado'] = ""
        info['nomVictima'] = ""
        info['nomSospechoso'] = ""
        info['fecha'] = time.strftime("%c")#raw_input("Ingrese la fhecha de la incuatacion: ")
        info['lugar'] = ""
        print("Generando archivo......")
        
        i = 0
        listaItems = []
        while (i<numEvidencias):
            item = {}
            item['numItem'] = str(i+1)
            item['nomArchivo'] = listaArchivos[i]
            item['desArchivo'] = subprocess.check_output('file '+ruta+listaArchivos[i],shell=True)
            item['hash'] = subprocess.check_output('md5sum '+ruta+listaArchivos[i],shell=True)[0:31]     
            listaItems.append(item)
            i = i+1
        info['archivos'] = listaItems
        generarWord(info)    

    else:
        print("Directorio vacio")
